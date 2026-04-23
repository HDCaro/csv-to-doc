import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn

BOOK_PATH = 'HITS AND HAPPINESS FINAL 2 Format.docx'
OUTPUT_FILE = 'Hits And Happiness Final 2 Discog.docx'
TABLE_ONLY_FILE = 'Richard Niles Discography Table.docx'

BOOKS_CSV = 'richard_niles_books.csv'


# ----------------------------
# Helpers
# ----------------------------
def set_cell_text(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT, hanging=False):
    cell.text = ""

    p = cell.paragraphs[0]
    p.alignment = align

    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1

    if hanging:
        pf.left_indent = Pt(8)
        pf.first_line_indent = Pt(-8)

    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = "Georgia"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Georgia')


def prevent_row_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:cantSplit'))


def keep_with_next(row):
    for cell in row.cells:
        for p in cell.paragraphs:
            p.paragraph_format.keep_with_next = True


def set_repeat_table_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)


def set_table_full_width(table):
    tbl = table._tbl
    tblPr = tbl.tblPr

    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), '5000')  # 100%

    tblPr.append(tblW)


def smart_text(text, max_len=40, hard_limit=70):
    text = str(text).strip()
    if len(text) <= max_len:
        return text
    if len(text) <= hard_limit:
        return text
    return text[:hard_limit - 3].rstrip() + "..."


def split_artists(text):
    parts = [p.strip() for p in str(text).split(",")]
    if len(parts) <= 1:
        return text
    return ",\n".join(parts)


def normalize_multiline(text):
    return str(text).replace("\\n", "\n")


# ----------------------------
# Discography Table
# ----------------------------
def compute_column_ratios(df_grouped):
    max_artist = max_album = max_details = 1

    for (_, artist, album), group in df_grouped:
        max_artist = max(max_artist, len(str(artist)))
        max_album = max(max_album, len(str(album)))

        details = group.iloc[0]['track_title'] if len(group) == 1 else f"{len(group)} tracks"
        max_details = max(max_details, len(details))

    total = max_artist + max_album + max_details + 6

    artist = (max_artist / total) * 0.85
    album = (max_album / total)
    details = (max_details / total) * 1.10
    role = (6 / total)

    total2 = artist + album + details + role

    return [artist/total2, album/total2, details/total2, role/total2]


def build_table(doc, grouped, col_ratios):

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.allow_autofit = False

    table.style = 'Table Grid'
    set_table_full_width(table)

    headers = ['Artist', 'Album', 'Details', 'Role']
    header_row = table.rows[0]

    for i, text in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_text(cell, text, WD_ALIGN_PARAGRAPH.CENTER)
        for run in cell.paragraphs[0].runs:
            run.bold = True

    set_repeat_table_header(header_row)

    current_year = None

    for (year, artist, album), group in grouped:

        if year != current_year:
            current_year = year
            row = table.add_row()
            merged = row.cells[0]
            for i in range(1, 4):
                merged = merged.merge(row.cells[i])

            p = merged.paragraphs[0]
            run = p.add_run(str(year))
            run.bold = True
            run.font.size = Pt(12)

            keep_with_next(row)

        roles = set()
        for _, r in group.iterrows():
            if str(r.get('producer')) == 'True': roles.add('P')
            if str(r.get('arranger')) == 'True': roles.add('A')
            if str(r.get('composer')) == 'True': roles.add('C')

        role_text = ', '.join(sorted(roles))

        if len(group) == 1:
            details = smart_text(group.iloc[0]['track_title'])
            album_label = "Single"
        else:
            details = f"{len(group)} tracks"
            album_label = smart_text(album)

        row = table.add_row()
        prevent_row_split(row)

        artist_text = split_artists(smart_text(artist, 60, 120))

        set_cell_text(row.cells[0], artist_text, hanging=True)
        set_cell_text(row.cells[1], album_label, hanging=True)
        set_cell_text(row.cells[2], details, hanging=True)
        set_cell_text(row.cells[3], role_text, WD_ALIGN_PARAGRAPH.CENTER)

    return table


# ----------------------------
# Books Table (NEW)
# ----------------------------
def build_books_table(doc):

    df = pd.read_csv(BOOKS_CSV, sep='\t', encoding='cp1252')
    df = df.fillna('')

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.allow_autofit = False

    table.style = 'Table Grid'
    set_table_full_width(table)

    headers = ['Year', 'Title', 'Author', 'Publisher']
    header_row = table.rows[0]

    for i, text in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_text(cell, text, WD_ALIGN_PARAGRAPH.CENTER)
        for run in cell.paragraphs[0].runs:
            run.bold = True

    set_repeat_table_header(header_row)

    for _, r in df.iterrows():

        row = table.add_row()
        prevent_row_split(row)

        set_cell_text(row.cells[0], str(r['year']))
        set_cell_text(row.cells[1], normalize_multiline(r['title']), hanging=True)
        set_cell_text(row.cells[2], normalize_multiline(r['author']), hanging=True)
        set_cell_text(row.cells[3], normalize_multiline(r['publisher']), hanging=True)

    return table


# ----------------------------
# Main
# ----------------------------
def create_discography():

    df = pd.read_csv('richard_niles_discography.csv', sep='\t', encoding='cp1252')
    df = df.fillna('')
    df['year'] = pd.to_numeric(df['year'], errors='coerce')
    df = df.dropna(subset=['year'])
    df['year'] = df['year'].astype(int)

    df_sorted = df.sort_values(['year', 'artist', 'album', 'track_title'])
    grouped = df_sorted.groupby(['year', 'artist', 'album'])

    col_ratios = compute_column_ratios(grouped)

    doc = Document(BOOK_PATH)
    doc.add_section(WD_SECTION.NEW_PAGE)

    # Title
    title = doc.add_paragraph("RICHARD NILES DISCOGRAPHY BY YEAR", style='Heading 1')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for run in title.runs:
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "Georgia"

    # Table
    build_table(doc, grouped, col_ratios)

    # 🔥 Books Section
    doc.add_page_break()

    books_title = doc.add_paragraph("BOOKS BY RICHARD NILES", style='Heading 1')
    books_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for run in books_title.runs:
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "Georgia"

    build_books_table(doc)

    doc.save(OUTPUT_FILE)

    print("\n✔ FINAL — discography + books table\n")


if __name__ == "__main__":
    create_discography()